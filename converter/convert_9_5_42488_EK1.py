# converter/convert_9_5_42488_EK1.py

import os
import csv
from docx import Document

INPUT_FILE = "data/doc/9.5.42488 EK (1).docx"
OUTPUT_FILE = "data/csv/9.5.42488_EK1.csv"

def convert_tabular_docx_to_csv(input_path, output_path):
    doc = Document(input_path)
    rows = []

    # İlk tabloyu bul
    if not doc.tables:
        print("❌ Belgede tablo bulunamadı.")
        return

    table = doc.tables[0]

    for row in table.rows[1:]:  # başlık satırını atla
        il_cell = row.cells[0].text.strip()
        konu_cell = row.cells[1].text.strip()

        # "Yatırım Konusu" hücresinde birden fazla satır varsa, ayır
        konular = konu_cell.split("\n")
        for konu in konular:
            konu = konu.strip()
            if konu:
                rows.append([il_cell, konu])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["İl", "Yatırım Konusu"])
        writer.writerows(rows)

    print(f"✅ CSV oluşturuldu: {output_path}")
    print(f"➡️ Toplam satır sayısı: {len(rows)}")

if __name__ == "__main__":
    convert_tabular_docx_to_csv(INPUT_FILE, OUTPUT_FILE)